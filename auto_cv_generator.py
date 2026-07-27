#!/usr/bin/env python3
"""
Auto CV Generator untuk JobStreet
Cara Pakai:
1. Jalankan: python3 auto_cv_generator.py
2. Pilih hasil scraping dari folder carilokermu
3. Isi data diri
4. CV otomatis dibuat dan disesuaikan dengan lowongan
"""

import os
import csv
import json
from datetime import datetime
from typing import List, Dict, Optional
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("❌ Error: python-docx belum terinstall!")
    print("   Jalankan: pip install python-docx")
    exit(1)

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.lib.units import cm, inch
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
except ImportError:
    print("❌ Error: reportlab belum terinstall!")
    print("   Jalankan: pip install reportlab")
    exit(1)

# Folder untuk menyimpan hasil
FOLDER_HASIL = Path(__file__).parent / "carilokermu"
FOLDER_CV = FOLDER_HASIL / "cv_generated"

# Pastikan folder ada
FOLDER_HASIL.mkdir(exist_ok=True)
FOLDER_CV.mkdir(exist_ok=True)


def tampilkan_file_scraping() -> List[Dict]:
    """Menampilkan semua file hasil scraping di folder carilokermu"""
    print("\n" + "="*60)
    print("📁 MENCARI FILE HASIL SCRAPING DI folder 'carilokermu'")
    print("="*60)
    
    file_list = []
    
    # Cari semua file CSV di folder
    for file_path in FOLDER_HASIL.glob("*.csv"):
        if file_path.name.startswith("~"):
            continue
            
        try:
            with open(file_path, 'r', encoding='utf-8-sig') as f:
                reader = csv.DictReader(f)
                rows = list(reader)
                if rows:
                    file_list.append({
                        'path': str(file_path),
                        'name': file_path.name,
                        'count': len(rows),
                        'data': rows
                    })
        except Exception as e:
            print(f"⚠️ Gagal membaca {file_path.name}: {e}")
    
    if not file_list:
        print("\n❌ Tidak ditemukan file hasil scraping di folder 'carilokermu'")
        print("   Silakan jalankan easy_search.py terlebih dahulu untuk mencari lowongan")
        return []
    
    # Tampilkan semua file
    for idx, file_info in enumerate(file_list, 1):
        print(f"\n[{idx}] 📄 {file_info['name']}")
        print(f"    Jumlah lowongan: {file_info['count']}")
        
        # Tampilkan 3 lowongan pertama sebagai preview
        if file_info['data']:
            print("    Preview:")
            for i, job in enumerate(file_info['data'][:3], 1):
                judul = job.get('judul', 'N/A')
                perusahaan = job.get('perusahaan', 'N/A')
                lokasi = job.get('lokasi', 'N/A')
                print(f"      {i}. {judul} - {perusahaan} ({lokasi})")
            if len(file_info['data']) > 3:
                print(f"      ... dan {len(file_info['data']) - 3} lowongan lainnya")
    
    return file_list


def pilih_file_scraping(file_list: List[Dict]) -> Optional[Dict]:
    """Meminta user memilih file scraping"""
    if not file_list:
        return None
    
    while True:
        try:
            pilihan = input(f"\nPilih file (1-{len(file_list)}): ")
            idx = int(pilihan) - 1
            if 0 <= idx < len(file_list):
                print(f"\n✅ File terpilih: {file_list[idx]['name']}")
                return file_list[idx]
            else:
                print(f"   ❌ Pilihan harus antara 1-{len(file_list)}")
        except ValueError:
            print("   ❌ Input harus angka")
        except KeyboardInterrupt:
            print("\n\n⚠️ Dibatalkan")
            return None


def tampilkan_lowongan(data: List[Dict], max_tampil: int = 20):
    """Menampilkan daftar lowongan dari file yang dipilih"""
    print("\n" + "="*60)
    print("📋 DAFTAR LOWONGAN TERSEDIA")
    print("="*60)
    
    total = len(data)
    to_show = min(total, max_tampil)
    
    for i, job in enumerate(data[:to_show], 1):
        no = job.get('no', i)
        judul = job.get('judul', 'N/A')
        perusahaan = job.get('perusahaan', 'N/A')
        lokasi = job.get('lokasi', 'N/A')
        gaji = job.get('gaji', '-')
        
        print(f"\n[{i}] {judul}")
        print(f"    🏢 {perusahaan}")
        print(f"    📍 {lokasi}")
        if gaji and gaji != '-':
            print(f"    💰 {gaji}")
    
    if total > max_tampil:
        print(f"\n   ... dan {total - max_tampil} lowongan lainnya")
    
    return to_show


def pilih_lowongan(data: List[Dict], max_tampil: int = 20) -> Optional[Dict]:
    """Meminta user memilih lowongan yang ingin dilamar"""
    while True:
        try:
            pilihan = input(f"\nPilih lowongan yang ingin dilamar (1-{max_tampil}): ")
            if pilihan.lower() == 'q':
                return None
            idx = int(pilihan) - 1
            if 0 <= idx < len(data):
                selected = data[idx]
                print(f"\n✅ Lowongan terpilih:")
                print(f"   📌 {selected.get('judul', 'N/A')}")
                print(f"   🏢 {selected.get('perusahaan', 'N/A')}")
                print(f"   📍 {selected.get('lokasi', 'N/A')}")
                return selected
            else:
                print(f"   ❌ Pilihan harus antara 1-{max_tampil}")
        except ValueError:
            print("   ❌ Input harus angka")
        except KeyboardInterrupt:
            print("\n\n⚠️ Dibatalkan")
            return None


def kumpulkan_data_diri() -> Dict:
    """Mengumpulkan data diri dari user"""
    print("\n" + "="*60)
    print("👤 ISI DATA DIRI ANDA")
    print("="*60)
    print("(Tekan Enter untuk melewati field opsional)")
    
    data = {}
    
    # Data pribadi wajib
    print("\n--- Data Pribadi ---")
    data['nama_lengkap'] = input("Nama Lengkap: ").strip()
    if not data['nama_lengkap']:
        print("   ❌ Nama lengkap wajib diisi!")
        return None
    
    data['email'] = input("Email: ").strip()
    data['telepon'] = input("Nomor Telepon/WA: ").strip()
    data['alamat'] = input("Alamat Lengkap: ").strip()
    data['kota'] = input("Kota: ").strip()
    data['tanggal_lahir'] = input("Tanggal Lahir (DD/MM/YYYY): ").strip()
    data['linkedin'] = input("LinkedIn URL (opsional): ").strip()
    data['portfolio'] = input("Portfolio/Website (opsional): ").strip()
    
    # Pendidikan
    print("\n--- Pendidikan Terakhir ---")
    data['pendidikan_nama'] = input("Nama Universitas/Sekolah: ").strip()
    data['pendidikan_jurusan'] = input("Jurusan: ").strip()
    data['pendidikan_gelar'] = input("Gelar (contoh: S.Kom, S.E): ").strip()
    data['pendidikan_tahun_masuk'] = input("Tahun Masuk: ").strip()
    data['pendidikan_tahun_lulus'] = input("Tahun Lulus: ").strip()
    data['pendidikan_ipk'] = input("IPK (opsional): ").strip()
    
    # Pengalaman Kerja
    print("\n--- Pengalaman Kerja (terakhir) ---")
    data['pengalaman_jumlah'] = input("Berapa banyak pengalaman kerja? (default 1): ").strip() or "1"
    
    pengalaman_list = []
    try:
        jumlah_pengalaman = int(data['pengalaman_jumlah'])
        for i in range(jumlah_pengalaman):
            print(f"\n  Pengalaman #{i+1}:")
            peng = {
                'posisi': input("    Posisi/Jabatan: ").strip(),
                'perusahaan': input("    Nama Perusahaan: ").strip(),
                'lokasi': input("    Lokasi: ").strip(),
                'tanggal_mulai': input("    Tanggal Mulai (MM/YYYY): ").strip(),
                'tanggal_selesai': input("    Tanggal Selesai (MM/YYYY atau 'sekarang'): ").strip(),
                'deskripsi': input("    Deskripsi pekerjaan (pisahkan dengan koma): ").strip()
            }
            pengalaman_list.append(peng)
    except ValueError:
        print("   ⚠️ Input tidak valid, menggunakan 1 pengalaman default")
        peng = {
            'posisi': input("  Posisi/Jabatan: ").strip(),
            'perusahaan': input("  Nama Perusahaan: ").strip(),
            'lokasi': input("  Lokasi: ").strip(),
            'tanggal_mulai': input("  Tanggal Mulai (MM/YYYY): ").strip(),
            'tanggal_selesai': input("  Tanggal Selesai (MM/YYYY atau 'sekarang'): ").strip(),
            'deskripsi': input("  Deskripsi pekerjaan: ").strip()
        }
        pengalaman_list = [peng]
    
    data['pengalaman'] = pengalaman_list
    
    # Keahlian
    print("\n--- Keahlian ---")
    print("(Pisahkan setiap keahlian dengan koma)")
    data['keahlian_teknis'] = input("Keahlian Teknis (contoh: Python, Excel, AutoCAD): ").strip()
    data['keahlian_soft_skill'] = input("Soft Skills (contoh: Komunikasi, Leadership): ").strip()
    data['bahasa'] = input("Bahasa (contoh: Indonesia (Native), Inggris (Fluent)): ").strip()
    
    # Sertifikasi (opsional)
    print("\n--- Sertifikasi (opsional) ---")
    data['sertifikasi'] = input("Sertifikasi (pisahkan dengan koma): ").strip()
    
    # Simpan data ke file JSON untuk penggunaan berikutnya
    file_data = FOLDER_HASIL / "data_diri.json"
    with open(file_data, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=2, ensure_ascii=False)
    
    print(f"\n✅ Data diri tersimpan di {file_data}")
    
    return data


def buat_cv_docx(data_diri: Dict, lowongan: Dict, nama_file: str):
    """Membuat CV dalam format DOCX"""
    doc = Document()
    
    # Style
    styles = doc.styles
    
    # Judul - Nama
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_name = title_para.add_run(data_diri['nama_lengkap'])
    run_name.bold = True
    run_name.font.size = Pt(18)
    
    # Kontak
    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    kontak_info = []
    if data_diri.get('email'):
        kontak_info.append(f"📧 {data_diri['email']}")
    if data_diri.get('telepon'):
        kontak_info.append(f"📱 {data_diri['telepon']}")
    if data_diri.get('alamat'):
        kontak_info.append(f"📍 {data_diri['alamat']}")
    if data_diri.get('linkedin'):
        kontak_info.append(f"💼 {data_diri['linkedin']}")
    
    contact_run = contact_para.add_run(" | ".join(kontak_info))
    contact_run.font.size = Pt(9)
    
    doc.add_paragraph()  # Spacing
    
    # Target Posisi
    target_para = doc.add_paragraph()
    target_run = target_para.add_run(f"Melamar sebagai: {lowongan.get('judul', 'Posisi yang Ditujukan')}")
    target_run.bold = True
    target_run.font.size = Pt(11)
    doc.add_paragraph(f"Perusahaan: {lowongan.get('perusahaan', 'N/A')} | Lokasi: {lowongan.get('lokasi', 'N/A')}")
    
    doc.add_paragraph()  # Spacing
    
    # Profil Singkat
    doc.add_heading('PROFIL SINGKAT', level=1)
    profil_text = f"Saya adalah profesional yang berpengalaman di bidang {data_diri.get('pendidikan_jurusan', 'umum')}. "
    profil_text += f"Dengan latar belakang pendidikan {data_diri.get('pendidikan_gelar', '')} dari {data_diri.get('pendidikan_nama', '')}, "
    profil_text += f"saya memiliki keahlian dalam {data_diri.get('keahlian_teknis', '')}. "
    profil_text += f"Saya tertarik untuk berkontribusi di {lowongan.get('perusahaan', 'perusahaan ini')} sebagai {lowongan.get('judul', 'posisi yang dituju')}."
    doc.add_paragraph(profil_text)
    
    # Pendidikan
    doc.add_heading('PENDIDIKAN', level=1)
    pend_table = doc.add_table(rows=1, cols=3)
    pend_table.style = 'Table Grid'
    hdr_cells = pend_table.rows[0].cells
    hdr_cells[0].text = 'Institusi'
    hdr_cells[1].text = 'Jurusan/Gelar'
    hdr_cells[2].text = 'Periode'
    
    for cell in hdr_cells:
        cell.paragraphs[0].runs[0].bold = True
    
    row_cells = pend_table.add_row().cells
    row_cells[0].text = data_diri.get('pendidikan_nama', '')
    row_cells[1].text = f"{data_diri.get('pendidikan_jurusan', '')} - {data_diri.get('pendidikan_gelar', '')}"
    periode = f"{data_diri.get('pendidikan_tahun_masuk', '')} - {data_diri.get('pendidikan_tahun_lulus', '')}"
    if data_diri.get('pendidikan_ipk'):
        periode += f" (IPK: {data_diri['pendidikan_ipk']})"
    row_cells[2].text = periode
    
    # Pengalaman Kerja
    doc.add_heading('PENGALAMAN KERJA', level=1)
    for peng in data_diri.get('pengalaman', []):
        p_posisi = doc.add_paragraph()
        run_pos = p_posisi.add_run(f"{peng.get('posisi', '')} - {peng.get('perusahaan', '')}")
        run_pos.bold = True
        doc.add_paragraph(f"📍 {peng.get('lokasi', '')} | {peng.get('tanggal_mulai', '')} - {peng.get('tanggal_selesai', '')}")
        
        deskripsi = peng.get('deskripsi', '')
        if deskripsi:
            for item in deskripsi.split(','):
                doc.add_paragraph(f"• {item.strip()}", style='List Bullet')
        doc.add_paragraph()  # Spacing antar pengalaman
    
    # Keahlian
    doc.add_heading('KEAHLIAN', level=1)
    if data_diri.get('keahlian_teknis'):
        doc.add_paragraph(f"Keahlian Teknis: {data_diri['keahlian_teknis']}")
    if data_diri.get('keahlian_soft_skill'):
        doc.add_paragraph(f"Soft Skills: {data_diri['keahlian_soft_skill']}")
    if data_diri.get('bahasa'):
        doc.add_paragraph(f"Bahasa: {data_diri['bahasa']}")
    
    # Sertifikasi
    if data_diri.get('sertifikasi'):
        doc.add_heading('SERTIFIKASI', level=1)
        for sert in data_diri['sertifikasi'].split(','):
            doc.add_paragraph(f"• {sert.strip()}", style='List Bullet')
    
    # Simpan file
    doc.save(nama_file)
    return nama_file


def buat_cv_pdf(data_diri: Dict, lowongan: Dict, nama_file: str):
    """Membuat CV dalam format PDF"""
    doc = SimpleDocTemplate(nama_file, pagesize=A4,
                           rightMargin=2*cm, leftMargin=2*cm,
                           topMargin=2*cm, bottomMargin=2*cm)
    
    elements = []
    styles = getSampleStyleSheet()
    
    # Custom styles
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=18,
        textColor=colors.HexColor('#2C3E50'),
        spaceAfter=12,
        alignment=1  # Center
    )
    
    heading_style = ParagraphStyle(
        'CustomHeading',
        parent=styles['Heading2'],
        fontSize=12,
        textColor=colors.HexColor('#34495E'),
        spaceBefore=12,
        spaceAfter=6,
        borderPadding=(3, 0, 3, 0),
        borderColor=colors.HexColor('#3498DB'),
        borderWidth=(0, 0, 2, 0)
    )
    
    normal_style = styles['Normal']
    normal_style.fontSize = 10
    
    # Nama
    elements.append(Paragraph(data_diri['nama_lengkap'], title_style))
    
    # Kontak
    kontak_info = []
    if data_diri.get('email'):
        kontak_info.append(f"📧 {data_diri['email']}")
    if data_diri.get('telepon'):
        kontak_info.append(f"📱 {data_diri['telepon']}")
    if data_diri.get('kota'):
        kontak_info.append(f"📍 {data_diri['kota']}")
    
    kontak_text = " | ".join(kontak_info)
    elements.append(Paragraph(kontak_text, normal_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Target Posisi
    target_text = f"<b>Melamar sebagai:</b> {lowongan.get('judul', 'Posisi yang Ditujuan')}<br/>"
    target_text += f"<b>Perusahaan:</b> {lowongan.get('perusahaan', 'N/A')} | <b>Lokasi:</b> {lowongan.get('lokasi', 'N/A')}"
    elements.append(Paragraph(target_text, normal_style))
    elements.append(Spacer(1, 0.5*cm))
    
    # Profil Singkat
    elements.append(Paragraph("PROFIL SINGKAT", heading_style))
    profil_text = f"Saya adalah profesional yang berpengalaman di bidang {data_diri.get('pendidikan_jurusan', 'umum')}. "
    profil_text += f"Dengan latar belakang pendidikan {data_diri.get('pendidikan_gelar', '')} dari {data_diri.get('pendidikan_nama', '')}, "
    profil_text += f"saya memiliki keahlian dalam {data_diri.get('keahlian_teknis', '')}. "
    profil_text += f"Saya tertarik untuk berkontribusi di {lowongan.get('perusahaan', 'perusahaan ini')} sebagai {lowongan.get('judul', 'posisi yang dituju')}."
    elements.append(Paragraph(profil_text, normal_style))
    elements.append(Spacer(1, 0.3*cm))
    
    # Pendidikan
    elements.append(Paragraph("PENDIDIKAN", heading_style))
    pend_data = [
        ['Institusi', 'Jurusan/Gelar', 'Periode'],
        [
            data_diri.get('pendidikan_nama', ''),
            f"{data_diri.get('pendidikan_jurusan', '')} - {data_diri.get('pendidikan_gelar', '')}",
            f"{data_diri.get('pendidikan_tahun_masuk', '')} - {data_diri.get('pendidikan_tahun_lulus', '')}"
        ]
    ]
    pend_table = Table(pend_data, colWidths=[5*cm, 7*cm, 4*cm])
    pend_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498DB')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 10),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 8),
        ('TOPPADDING', (0, 0), (-1, 0), 8),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(pend_table)
    elements.append(Spacer(1, 0.3*cm))
    
    # Pengalaman Kerja
    elements.append(Paragraph("PENGALAMAN KERJA", heading_style))
    for peng in data_diri.get('pengalaman', []):
        elements.append(Paragraph(f"<b>{peng.get('posisi', '')}</b> - {peng.get('perusahaan', '')}", normal_style))
        elements.append(Paragraph(f"📍 {peng.get('lokasi', '')} | {peng.get('tanggal_mulai', '')} - {peng.get('tanggal_selesai', '')}", normal_style))
        deskripsi = peng.get('deskripsi', '')
        if deskripsi:
            for item in deskripsi.split(','):
                elements.append(Paragraph(f"• {item.strip()}", normal_style))
        elements.append(Spacer(1, 0.2*cm))
    
    # Keahlian
    elements.append(Paragraph("KEAHLIAN", heading_style))
    if data_diri.get('keahlian_teknis'):
        elements.append(Paragraph(f"<b>Keahlian Teknis:</b> {data_diri['keahlian_teknis']}", normal_style))
    if data_diri.get('keahlian_soft_skill'):
        elements.append(Paragraph(f"<b>Soft Skills:</b> {data_diri['keahlian_soft_skill']}", normal_style))
    if data_diri.get('bahasa'):
        elements.append(Paragraph(f"<b>Bahasa:</b> {data_diri['bahasa']}", normal_style))
    
    # Sertifikasi
    if data_diri.get('sertifikasi'):
        elements.append(Paragraph("SERTIFIKASI", heading_style))
        for sert in data_diri['sertifikasi'].split(','):
            elements.append(Paragraph(f"• {sert.strip()}", normal_style))
    
    # Build PDF
    doc.build(elements)
    return nama_file


def buat_cover_letter(data_diri: Dict, lowongan: Dict, nama_file: str):
    """Membuat surat lamaran kerja (Cover Letter)"""
    doc = Document()
    
    # Header
    header_para = doc.add_paragraph()
    header_run = header_para.add_run(data_diri['nama_lengkap'])
    header_run.bold = True
    
    kontak_info = []
    if data_diri.get('email'):
        kontak_info.append(data_diri['email'])
    if data_diri.get('telepon'):
        kontak_info.append(data_diri['telepon'])
    
    if kontak_info:
        doc.add_paragraph(", ".join(kontak_info))
    doc.add_paragraph()
    
    # Tanggal
    tanggal = datetime.now().strftime('%d %B %Y')
    doc.add_paragraph(tanggal)
    doc.add_paragraph()
    
    # Tujuan
    doc.add_paragraph(f"Kepada Yth.,")
    doc.add_paragraph(f"HRD {lowongan.get('perusahaan', 'Perusahaan')}")
    doc.add_paragraph(f"Di Tempat")
    doc.add_paragraph()
    
    # Salam pembuka
    doc.add_paragraph("Dengan hormat,")
    doc.add_paragraph()
    
    # Paragraf pembuka
    doc.add_paragraph("Saya yang bertanda tangan di bawah ini:")
    doc.add_paragraph()
    
    # Data pelamar - menggunakan list bukan table
    doc.add_paragraph(f"Nama          : {data_diri['nama_lengkap']}")
    doc.add_paragraph(f"Email         : {data_diri.get('email', '-')}")
    doc.add_paragraph(f"Telepon       : {data_diri.get('telepon', '-')}")
    doc.add_paragraph(f"Lamaran       : {lowongan.get('judul', 'Posisi')} di {lowongan.get('perusahaan', 'Perusahaan')}")
    doc.add_paragraph()
    
    # Paragraf isi
    pengalaman_tahun = ""
    if data_diri.get('pengalaman') and len(data_diri['pengalaman']) > 0:
        peng_first = data_diri['pengalaman'][0]
        tanggal_mulai = peng_first.get('tanggal_mulai', '')
        if '/' in tanggal_mulai:
            tahun_mulai = tanggal_mulai.split('/')[-1]
            try:
                tahun_sekarang = datetime.now().year
                pengalaman_tahun = f"{tahun_sekarang - int(tahun_mulai)}"
            except:
                pengalaman_tahun = "1"
        else:
            pengalaman_tahun = "1"
    
    isi = f"""Dengan ini saya bermaksud melamar pekerjaan di {lowongan.get('perusahaan', 'perusahaan Anda')} untuk posisi {lowongan.get('judul', 'posisi yang dituju')} sebagaimana diinformasikan melalui JobStreet.

Saya merupakan lulusan {data_diri.get('pendidikan_gelar', '')} {data_diri.get('pendidikan_jurusan', '')} dari {data_diri.get('pendidikan_nama', '')}. Saya memiliki pengalaman kerja sebagai {data_diri.get('pengalaman', [{}])[0].get('posisi', 'profesional')} di {data_diri.get('pengalaman', [{}])[0].get('perusahaan', 'perusahaan sebelumnya')} selama {pengalaman_tahun} tahun.

Keahlian utama saya meliputi {data_diri.get('keahlian_teknis', '')}. Saya yakin dengan kualifikasi dan pengalaman yang saya miliki, saya dapat memberikan kontribusi yang positif bagi {lowongan.get('perusahaan', 'perusahaan Anda')}.

Sebagai bahan pertimbangan, bersama ini saya lampirkan Curriculum Vitae (CV) dan dokumen pendukung lainnya."""
    
    doc.add_paragraph(isi)
    doc.add_paragraph()
    
    # Penutup
    penutup = """Demikian surat lamaran ini saya buat dengan sebenar-benarnya. Besar harapan saya untuk dapat mengikuti tahap seleksi selanjutnya. Atas perhatian dan kesempatan yang diberikan, saya ucapkan terima kasih."""
    doc.add_paragraph(penutup)
    doc.add_paragraph()
    doc.add_paragraph("Hormat saya,")
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()
    sign_para = doc.add_paragraph()
    sign_run = sign_para.add_run(data_diri['nama_lengkap'])
    sign_run.bold = True
    
    # Simpan
    doc.save(nama_file)
    return nama_file


def main():
    """Fungsi utama"""
    print("\n" + "="*60)
    print("🎯 AUTO CV GENERATOR - JOBSTREET")
    print("="*60)
    
    # Langkah 1: Tampilkan file scraping
    file_list = tampilkan_file_scraping()
    if not file_list:
        return
    
    # Langkah 2: Pilih file scraping
    selected_file = pilih_file_scraping(file_list)
    if not selected_file:
        return
    
    # Langkah 3: Tampilkan dan pilih lowongan
    tampilkan_lowongan(selected_file['data'])
    lowongan = pilih_lowongan(selected_file['data'])
    if not lowongan:
        return
    
    # Langkah 4: Kumpulkan data diri
    data_diri = kumpulkan_data_diri()
    if not data_diri:
        print("\n❌ Data diri tidak lengkap")
        return
    
    # Langkah 5: Buat CV
    print("\n" + "="*60)
    print("📝 MEMBUAT CV DAN SURAT LAMARAN...")
    print("="*60)
    
    # Buat nama file
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    nama_perusahaan = lowongan.get('perusahaan', 'perusahaan').replace(' ', '_').replace('/', '_')
    nama_posisi = lowongan.get('judul', 'posisi').replace(' ', '_').replace('/', '_')[:30]
    
    base_filename = f"CV_{data_diri['nama_lengkap'].replace(' ', '_')}_{nama_posisi}_{timestamp}"
    
    # Buat DOCX
    file_docx = str(FOLDER_CV / f"{base_filename}.docx")
    try:
        buat_cv_docx(data_diri, lowongan, file_docx)
        print(f"✅ CV DOCX: {file_docx}")
    except Exception as e:
        print(f"❌ Gagal membuat DOCX: {e}")
    
    # Buat PDF
    file_pdf = str(FOLDER_CV / f"{base_filename}.pdf")
    try:
        buat_cv_pdf(data_diri, lowongan, file_pdf)
        print(f"✅ CV PDF: {file_pdf}")
    except Exception as e:
        print(f"❌ Gagal membuat PDF: {e}")
    
    # Buat Cover Letter
    file_cl = str(FOLDER_CV / f"CoverLetter_{data_diri['nama_lengkap'].replace(' ', '_')}_{nama_posisi}_{timestamp}.docx")
    try:
        buat_cover_letter(data_diri, lowongan, file_cl)
        print(f"✅ Cover Letter: {file_cl}")
    except Exception as e:
        print(f"❌ Gagal membuat Cover Letter: {e}")
    
    # Ringkasan
    print("\n" + "="*60)
    print("🎉 SELESAI!")
    print("="*60)
    print(f"\n📌 Lowongan yang dilamar:")
    print(f"   • Posisi: {lowongan.get('judul', 'N/A')}")
    print(f"   • Perusahaan: {lowongan.get('perusahaan', 'N/A')}")
    print(f"   • Lokasi: {lowongan.get('lokasi', 'N/A')}")
    print(f"   • Link: {lowongan.get('link', 'N/A')}")
    print(f"\n📁 File tersimpan di folder: {FOLDER_CV}")
    print("\n💡 Tips:")
    print("   1. Review CV sebelum mengirim")
    print("   2. Sesuaikan lagi jika perlu")
    print("   3. Upload CV ke JobStreet dan apply langsung")
    print("="*60 + "\n")


if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n⚠️ Program dihentikan oleh user")
    except Exception as e:
        print(f"\n❌ Error tak terduga: {e}")
        import traceback
        traceback.print_exc()
