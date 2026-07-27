#!/usr/bin/env python3
"""
Auto CV Generator - Pilih Lowongan & Generate CV Otomatis
Membaca hasil scraping, tampilkan pilihan, generate CV sesuai posisi
"""

import json
import csv
import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import subprocess

# Konfigurasi folder
FOLDER_UTAMA = os.path.dirname(os.path.abspath(__file__))
FOLDER_LOKER = os.path.join(FOLDER_UTAMA, 'carilokermu')
FILE_DATA_DIRI = os.path.join(FOLDER_LOKER, 'data_diri.json')
FOLDER_CV_OUTPUT = os.path.join(FOLDER_LOKER, 'cv_generated')

def load_data_diri():
    """Load data diri dari JSON"""
    if not os.path.exists(FILE_DATA_DIRI):
        print("❌ File data_diri.json tidak ditemukan!")
        return None

    with open(FILE_DATA_DIRI, 'r', encoding='utf-8') as f:
        return json.load(f)

def load_lowongan():
    """Load semua lowongan dari file CSV terbaru"""
    csv_files = [f for f in os.listdir(FOLDER_LOKER) if f.startswith('loker_') and f.endswith('.csv')]

    if not csv_files:
        print("❌ Tidak ada file hasil scraping ditemukan!")
        return []

    # Ambil file terbaru
    csv_files.sort(reverse=True)
    file_terbaru = os.path.join(FOLDER_LOKER, csv_files[0])

    print(f"📄 Membaca dari: {csv_files[0]}")

    lowongan = []
    with open(file_terbaru, 'r', encoding='utf-8') as f:
        reader = csv.DictReader(f)
        for row in reader:
            # Support multiple column names for job title
            judul = row.get('Posisi') or row.get('judul') or row.get('judul_lowongan') or row.get('title') or row.get('posisi')
            if judul:  # Skip baris kosong
                lowongan.append(row)

    return lowongan

def get_field(loker, field_name, default=''):
    """Helper untuk mengambil field dengan support multiple column names"""
    # Mapping field names yang mungkin
    field_mappings = {
        'judul': ['Posisi', 'judul', 'judul_lowongan', 'title', 'posisi', 'position'],
        'perusahaan': ['Perusahaan', 'perusahaan', 'company', 'nama_perusahaan'],
        'lokasi': ['Lokasi', 'lokasi', 'location', 'kota'],
        'gaji': ['gaji', 'salary', 'rentang_gaji'],
        'tanggal_scrape': ['tanggal_scrape', 'tanggal_posting', 'posted_date', 'tanggal'],
        'link': ['Link', 'link', 'url', 'website', 'job_url']
    }
    
    if field_name in field_mappings:
        for key in field_mappings[field_name]:
            if key in loker and loker[key]:
                return loker[key]
    return default

def tampilkan_lowongan(lowongan):
    """Tampilkan daftar lowongan dengan format rapi"""
    print("\n" + "="*80)
    print("📋 DAFTAR LOWONGAN TERSEDIA")
    print("="*80)

    if not lowongan:
        print("Tidak ada lowongan tersedia.")
        return

    for i, loker in enumerate(lowongan, 1):
        print(f"\n[{i}] {get_field(loker, 'judul', 'N/A')}")
        print(f"    🏢 Perusahaan: {get_field(loker, 'perusahaan', 'N/A')}")
        print(f"    📍 Lokasi: {get_field(loker, 'lokasi', 'N/A')}")
        print(f"    💰 Gaji: {get_field(loker, 'gaji', 'N/A')}")
        print(f"    📅 Scraped: {get_field(loker, 'tanggal_scrape', 'N/A')}")
        print(f"    🔗 Link: {get_field(loker, 'link', 'N/A')}")

    print("\n" + "="*80)

def pilih_lowongan(lowongan):
    """Minta user memilih lowongan"""
    while True:
        try:
            pilihan = input("\n👉 Masukkan nomor lowongan yang ingin dilamar (atau 0 untuk keluar): ")
            pilihan = int(pilihan)

            if pilihan == 0:
                return None
            elif 1 <= pilihan <= len(lowongan):
                return lowongan[pilihan - 1]
            else:
                print(f"⚠️  Pilihan harus antara 1-{len(lowongan)} atau 0 untuk keluar")
        except ValueError:
            print("⚠️  Masukkan angka yang valid!")

def generate_cv(data_diri, lowongan):
    """Generate CV dalam format DOCX dan PDF"""
    if not os.path.exists(FOLDER_CV_OUTPUT):
        os.makedirs(FOLDER_CV_OUTPUT)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    nama_file_safe = "".join(c for c in lowongan['judul'] if c.isalnum() or c in (' ', '-', '_')).strip()
    nama_file_safe = nama_file_safe.replace(' ', '_')

    # Nama file
    nama_depan = data_diri['nama_lengkap'].split()[0]
    nama_belakang = data_diri['nama_lengkap'].split()[-1] if len(data_diri['nama_lengkap'].split()) > 1 else ''

    judul_lowongan = get_field(lowongan, 'judul', 'Lowongan')
    nama_file_safe = "".join(c for c in judul_lowongan if c.isalnum() or c in (' ', '-', '_')).strip()
    nama_file_safe = nama_file_safe.replace(' ', '_')

    filename_base = f"CV_{nama_depan}_{nama_belakang}_{nama_file_safe}_{timestamp}"
    file_docx = os.path.join(FOLDER_CV_OUTPUT, f"{filename_base}.docx")
    file_pdf = os.path.join(FOLDER_CV_OUTPUT, f"{filename_base}.pdf")

    # Buat dokumen Word
    doc = Document()

    # Header - Nama
    header = doc.add_heading(data_diri['nama_lengkap'], 0)
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Kontak
    kontak = f"{data_diri['email']} | {data_diri['telepon']} | {data_diri['alamat']}, {data_diri['kota']}"
    if data_diri.get('linkedin'):
        kontak += f" | LinkedIn: {data_diri['linkedin']}"

    p_kontak = doc.add_paragraph(kontak)
    p_kontak.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Posisi yang dilamar
    doc.add_heading(f"Melamar sebagai: {get_field(lowongan, 'judul', 'Posisi')}", level=1)
    doc.add_paragraph(f"Perusahaan: {get_field(lowongan, 'perusahaan', 'N/A')} | Lokasi: {get_field(lowongan, 'lokasi', 'N/A')}")

    # Profil Singkat
    doc.add_heading('Profil Singkat', level=2)
    profil = f"Lulusan {data_diri['pendidikan_nama']} jurusan {data_diri['pendidikan_jurusan']} dengan IPK {data_diri['pendidikan_ipk']}. "
    profil += f"Berpengalaman {data_diri['pengalaman_jumlah']} tahun di bidang teknologi dengan keahlian {data_diri['keahlian_teknis']}. "
    profil += "Memiliki kemampuan komunikasi yang baik dan mampu bekerja dalam tim maupun individu."
    doc.add_paragraph(profil)

    # Pendidikan
    doc.add_heading('Pendidikan', level=2)
    pend = doc.add_paragraph()
    pend.add_run(f"{data_diri['pendidikan_nama']}").bold = True
    pend.add_run(f"\n{data_diri['pendidikan_jurusan']} ({data_diri['pendidikan_gelar']})")
    pend.add_run(f"\n{data_diri['pendidikan_tahun_masuk']} - {data_diri['pendidikan_tahun_lulus']}")
    pend.add_run(f"\nIPK: {data_diri['pendidikan_ipk']}")

    # Pengalaman Kerja
    doc.add_heading('Pengalaman Kerja', level=2)
    for pengalaman in data_diri.get('pengalaman', []):
        p_exp = doc.add_paragraph()
        p_exp.add_run(f"{pengalaman['posisi']}").bold = True
        p_exp.add_run(f"\n{pengalaman['perusahaan']} | {pengalaman['lokasi']}")
        p_exp.add_run(f"\n{pengalaman['tanggal_mulai']} - {pengalaman['tanggal_selesai']}")

        # Deskripsi tugas
        deskripsi_list = pengalaman['deskripsi'].split(',')
        for tugas in deskripsi_list:
            doc.add_paragraph(tugas.strip(), style='List Bullet')

    # Keahlian
    doc.add_heading('Keahlian', level=2)
    doc.add_paragraph(f"Teknis: {data_diri['keahlian_teknis']}")
    doc.add_paragraph(f"Soft Skills: {data_diri['keahlian_soft_skill']}")
    doc.add_paragraph(f"Bahasa: {data_diri['bahasa']}")

    # Sertifikasi
    if data_diri.get('sertifikasi'):
        doc.add_heading('Sertifikasi', level=2)
        doc.add_paragraph(data_diri['sertifikasi'])

    # Simpan DOCX
    doc.save(file_docx)
    print(f"✅ CV berhasil dibuat: {os.path.basename(file_docx)}")

    # Convert ke PDF (jika libreoffice tersedia)
    try:
        subprocess.run([
            'libreoffice', '--headless', '--convert-to', 'pdf',
            '--outdir', FOLDER_CV_OUTPUT, file_docx
        ], check=True, capture_output=True)
        print(f"✅ PDF berhasil dibuat: {os.path.basename(file_pdf)}")
    except Exception as e:
        print(f"⚠️  Konversi PDF gagal (LibreOffice tidak tersedia): {e}")
        print("   File DOCX tetap dapat digunakan")

    return file_docx

def generate_cover_letter(data_diri, lowongan):
    """Generate Cover Letter"""
    if not os.path.exists(FOLDER_CV_OUTPUT):
        os.makedirs(FOLDER_CV_OUTPUT)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    judul_lowongan = get_field(lowongan, 'judul', 'Lowongan')
    nama_file_safe = "".join(c for c in judul_lowongan if c.isalnum() or c in (' ', '-', '_')).strip()
    nama_file_safe = nama_file_safe.replace(' ', '_')

    nama_depan = data_diri['nama_lengkap'].split()[0]
    nama_belakang = data_diri['nama_lengkap'].split()[-1] if len(data_diri['nama_lengkap'].split()) > 1 else ''

    filename_base = f"CoverLetter_{nama_depan}_{nama_belakang}_{nama_file_safe}_{timestamp}"
    file_docx = os.path.join(FOLDER_CV_OUTPUT, f"{filename_base}.docx")

    doc = Document()

    # Header
    doc.add_heading(data_diri['nama_lengkap'], level=1)
    doc.add_paragraph(f"{data_diri['email']} | {data_diri['telepon']} | {data_diri['kota']}")

    doc.add_paragraph()

    # Tanggal
    tanggal = datetime.now().strftime("%d %B %Y")
    doc.add_paragraph(tanggal)

    # Alamat perusahaan
    doc.add_paragraph(f"HRD Department\n{get_field(lowongan, 'perusahaan', 'Perusahaan')}\n{get_field(lowongan, 'lokasi', 'Lokasi')}")

    doc.add_paragraph()

    # Salam pembuka
    doc.add_paragraph("Dengan hormat,", style='Normal')

    # Isi cover letter
    isi = f"""Saya menulis surat ini untuk menyampaikan ketertarikan saya pada posisi {get_field(lowongan, 'judul', 'Posisi')} di {get_field(lowongan, 'perusahaan', 'Perusahaan')} sebagaimana diinformasikan melalui JobStreet.

Sebagai lulusan {data_diri['pendidikan_nama']} jurusan {data_diri['pendidikan_jurusan']} dengan pengalaman {data_diri['pengalaman_jumlah']} tahun di bidang teknologi, saya yakin memiliki kualifikasi yang sesuai dengan kebutuhan perusahaan Bapak/Ibu.

Selama karir saya, saya telah mengembangkan keahlian dalam {data_diri['keahlian_teknis']}, serta memiliki soft skills yang kuat dalam {data_diri['keahlian_soft_skill']}. Saya juga telah menyelesaikan berbagai sertifikasi profesional termasuk {data_diri.get('sertifikasi', 'berbagai pelatihan')}.

Saya sangat antusias untuk dapat berkontribusi di {get_field(lowongan, 'perusahaan', 'Perusahaan')} dan yakin bahwa pengalaman serta kemampuan saya dapat memberikan nilai tambah bagi perusahaan.

Bersama surat ini, saya lampirkan Curriculum Vitae untuk memberikan gambaran lebih detail mengenai kualifikasi saya. Saya sangat mengharapkan kesempatan untuk dapat diskusi lebih lanjut mengenai bagaimana saya dapat berkontribusi pada tim Bapak/Ibu.

Terima kasih atas waktu dan pertimbangan Bapak/Ibu."""

    for paragraf in isi.split('\n\n'):
        doc.add_paragraph(paragraf.strip())

    # Penutup
    doc.add_paragraph()
    doc.add_paragraph("Hormat saya,")
    doc.add_paragraph()
    doc.add_paragraph(data_diri['nama_lengkap'])

    doc.save(file_docx)
    print(f"✅ Cover Letter berhasil dibuat: {os.path.basename(file_docx)}")

    return file_docx

def main():
    print("="*80)
    print("🚀 AUTO CV GENERATOR - PILIH LOWONGAN & LAMAR OTOMATIS")
    print("="*80)

    # Load data diri
    print("\n📂 Memuat data diri...")
    data_diri = load_data_diri()
    if not data_diri:
        print("❌ Gagal memuat data diri. Pastikan file data_diri.json sudah diisi.")
        return

    print(f"✅ Data diri dimuat: {data_diri['nama_lengkap']}")

    # Load lowongan
    print("\n📂 Memuat hasil scraping...")
    lowongan = load_lowongan()

    if not lowongan:
        print("❌ Tidak ada lowongan tersedia. Jalankan easy_search.py terlebih dahulu.")
        return

    # Tampilkan pilihan
    tampilkan_lowongan(lowongan)

    # Pilih lowongan
    lowongan_dipilih = pilih_lowongan(lowongan)

    if not lowongan_dipilih:
        print("\n👋 Program dihentikan.")
        return

    print(f"\n🎯 Anda memilih: {get_field(lowongan_dipilih, 'judul', 'N/A')} di {get_field(lowongan_dipilih, 'perusahaan', 'N/A')}")

    # Konfirmasi
    konfirmasi = input("\n👉 Lanjutkan generate CV & Cover Letter? (y/n): ").lower()
    if konfirmasi != 'y':
        print("\n👋 Program dihentikan.")
        return

    # Generate CV
    print("\n⚙️  Generating CV...")
    cv_file = generate_cv(data_diri, lowongan_dipilih)

    # Generate Cover Letter
    print("\n⚙️  Generating Cover Letter...")
    cl_file = generate_cover_letter(data_diri, lowongan_dipilih)

    # Summary
    print("\n" + "="*80)
    print("✅ PROSES SELESAI!")
    print("="*80)
    print(f"\n📄 Files generated:")
    print(f"   • CV: {cv_file}")
    print(f"   • Cover Letter: {cl_file}")
    print(f"\n📁 Lokasi: {FOLDER_CV_OUTPUT}")
    print(f"\n💡 Tips:")
    print(f"   • Review dokumen sebelum dikirim")
    print(f"   • Sesuaikan jika ada informasi spesifik dari perusahaan")
    print(f"   • Lamar melalui link: {get_field(lowongan_dipilih, 'link', '#')}")
    print("\n" + "="*80)

if __name__ == "__main__":
    main()
